import os
import uuid
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def generate_word(analysis: str, transcript: str = None) -> str:
    """
    Генерация Word протокола встречи
    
    Args:
        analysis: Текст анализа от GPT
        transcript: Полный транскрипт (опционально)
    
    Returns:
        Путь к созданному файлу
    """
    try:
        doc = Document()
        
        # === Заголовок ===
        title = doc.add_heading('ПРОТОКОЛ ВСТРЕЧИ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Дата
        date_para = doc.add_paragraph()
        date_para.add_run(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M')}").bold = True
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Пустая строка
        
        # === Summary ===
        doc.add_heading('1. Резюме встречи', 1)
        summary = extract_section(analysis, "Summary")
        doc.add_paragraph(summary)
        
        # === Задачи ===
        doc.add_heading('2. Действия и задачи', 1)
        
        tasks = parse_tasks_from_analysis(analysis)
        
        if tasks:
            # Создаём таблицу
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Заголовки
            headers = table.rows[0].cells
            headers[0].text = 'Задача'
            headers[1].text = 'Ответственный'
            headers[2].text = 'Дедлайн'
            headers[3].text = 'Приоритет'
            
            # Делаем заголовки жирными
            for cell in headers:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Добавляем задачи
            for task in tasks:
                row = table.add_row().cells
                row[0].text = task.get('description', '')
                row[1].text = task.get('assignee', 'Не указан')
                row[2].text = task.get('deadline', 'Не указан')
                row[3].text = task.get('priority', 'Средний')
        else:
            doc.add_paragraph("Задачи не обнаружены в транскрипте.")
        
        # === Полный анализ ===
        doc.add_heading('3. Детальный анализ', 1)
        
        # Разбиваем анализ на секции
        sections = split_analysis_sections(analysis)
        for section_title, section_text in sections.items():
            if section_title not in ["Summary", "Задачи"]:
                doc.add_heading(section_title, 2)
                doc.add_paragraph(section_text)
        
        # === Транскрипт ===
        if transcript:
            doc.add_page_break()
            doc.add_heading('4. Полный транскрипт', 1)
            
            # Добавляем транскрипт меньшим шрифтом
            transcript_para = doc.add_paragraph(transcript)
            for run in transcript_para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(80, 80, 80)
        
        # === Футер ===
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_para.add_run("\n" + "─" * 60 + "\n").italic = True
        footer_para.add_run("Документ создан автоматически сервисом Voice2Action\n").italic = True
        footer_para.add_run(f"Дата создания: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Сохранение файла
        file_id = str(uuid.uuid4())
        filename = f"protocol_{file_id}.docx"
        file_path = f"/tmp/{filename}"
        
        doc.save(file_path)
        logger.info(f"Word file created: {file_path}")
        
        return file_path
    
    except Exception as e:
        logger.error(f"Word generation failed: {e}")
        raise Exception(f"Не удалось создать Word файл: {str(e)}")


def extract_section(text: str, section_name: str) -> str:
    """Извлечение секции из текста"""
    lines = text.split('\n')
    result = []
    in_section = False
    
    for line in lines:
        if section_name.lower() in line.lower() and '#' in line:
            in_section = True
            continue
        if in_section:
            if line.strip().startswith('##'):
                break
            if line.strip():
                result.append(line.strip())
    
    return '\n'.join(result) if result else "Не найдено"


def split_analysis_sections(analysis: str) -> dict:
    """Разбить анализ на секции по заголовкам ##"""
    sections = {}
    lines = analysis.split('\n')
    
    current_section = None
    current_text = []
    
    for line in lines:
        if line.strip().startswith('##'):
            if current_section:
                sections[current_section] = '\n'.join(current_text)
            current_section = line.replace('##', '').strip()
            current_text = []
        elif current_section:
            if line.strip():
                current_text.append(line.strip())
    
    if current_section:
        sections[current_section] = '\n'.join(current_text)
    
    return sections


def parse_tasks_from_analysis(analysis: str) -> list:
    """Парсинг задач из анализа"""
    tasks = []
    lines = analysis.split('\n')
    
    current_task = None
    for line in lines:
        if 'Задача:' in line:
            if current_task:
                tasks.append(current_task)
            
            current_task = {
                'description': line.split('Задача:')[1].strip(),
                'deadline': 'Не указан',
                'assignee': 'Не указан',
                'priority': 'Средний'
            }
        elif current_task:
            if 'Дедлайн:' in line:
                current_task['deadline'] = line.split('Дедлайн:')[1].strip()
            elif 'Ответственный:' in line:
                current_task['assignee'] = line.split('Ответственный:')[1].strip()
            elif 'Приоритет:' in line:
                current_task['priority'] = line.split('Приоритет:')[1].strip()
    
    if current_task:
        tasks.append(current_task)
    
    return tasks
